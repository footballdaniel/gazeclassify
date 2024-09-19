from matplotlib import pyplot as plt
import numpy as np
import pandas as pd
from sklearn.metrics import ConfusionMatrixDisplay, classification_report

from .persistence import SaveData, SaveFigure, SaveTable

from typing import List
import docx

class Table:

    def __init__(self):
        self.doc = docx.Document()
        run = self.doc.add_paragraph().add_run()
        run.font.size = docx.shared.Pt(12)
        run.font.name = 'Times New Roman'
        self.rows = []
        self.columns = 0

    def add_row(self, row: List[str]):
        self.rows.append(row)
        return self

    def add_columns(self, columns: List[str]):
        self.columns = columns
        return self

    def _columns_to(self, table):
        columns = table.rows[0].cells
        for i in range(len(self.columns)):
            columns[i].text = self.columns[i]

    def _rows_to(self, table):
        for row in range(len(self.rows)):
            for column in range(len(self.columns)):
                content = self.rows[row][column]

                # Sanitize
                if isinstance(content, float):
                    content = "{:.2f}".format(content)
                if not isinstance(content, str):
                    content = str(content)

                # Assign
                table.rows[row+1].cells[column].text = content

    def build(self):
        table = self.doc.add_table(rows=len(self.rows)+1, cols=len(self.columns))
        self._columns_to(table)
        self._rows_to(table)
        return self.doc




class  ClassificationScore:

    def __init__(self, persistence, logger):
        self.persistence = persistence
        self.logger = logger


    def f1_metrics(self, df_long, desired_order: List[str]):
        df = df_long.pivot(index=['Trial', 'Frame'], columns='Rater', values='Label')
        df = df[~(df.isna()).any(axis=1)]

        # Get all columns starting with Rater
        df['HumanRatersAgreeingCount'] = df.loc[:, df.columns.str.startswith("Rater")].apply(lambda x: x.value_counts().max(), axis=1)
        df['MajorityVote'] = df.loc[:, df.columns.str.startswith("Rater")].apply(lambda x: x.value_counts().idxmax(), axis=1)
        df["AtLeastTwoRatersAgree"] = df["HumanRatersAgreeingCount"] >= 2
        df = df[df["AtLeastTwoRatersAgree"]]

        y_true = df.MajorityVote.values
        y_pred = df.Algorithm.values
                
        # F1 score
        # https://www.youtube.com/watch?v=_dWNReULL9E
        report = classification_report(y_true, y_pred, output_dict=True, zero_division=0)        
        df_report = pd.DataFrame.from_dict(report).T

        rows = df_report.reset_index().values.tolist()
        headers = df_report.columns.str.capitalize().tolist()

        headers.insert(0, "AOI")
        headers = [col for col in headers if col != "Support"]
        rows = [row for row in rows if row[0] != "accuracy"]
        rows = [row for row in rows if row[0] != "macro avg"]
        rows = [row for row in rows if row[0] != "weighted avg"]

        # order the rows:
        rows_dict = {row[0]: row for row in rows}
        ordered_rows = [rows_dict[item] for item in desired_order if item in rows_dict]

        report_text = classification_report(y_true, y_pred, zero_division=0) # can be printed

        # get weighted avg
        weighted_avg = df_report.loc["weighted avg"].values.tolist()

        self.logger(f"Weighted F1 score: {weighted_avg[2]:.2f}")
        self.logger(f"Weighted precision: {weighted_avg[0]:.2f}")
        self.logger(f"Weighted recall: {weighted_avg[1]:.2f}")

        # Save weighted f1 score
        self.persistence.add_result(SaveData(f"{weighted_avg[2]:.2f}", "Weighted F1 score"))
        self.persistence.add_result(SaveData(f"{weighted_avg[0]:.2f}", "Weighted precision"))
        self.persistence.add_result(SaveData(f"{weighted_avg[1]:.2f}", "Weighted recall"))

        table = Table()
        table.add_columns(headers)
        for row in ordered_rows:
            table.add_row(row)
        doc = table.build()
        self.persistence.add_table(SaveTable(doc, "f1_metrics"))

    
    def confusion_matrix(self, df_long, desired_order: List[str]):
        df = df_long.pivot(index=['Trial', 'Frame'], columns='Rater', values='Label')
        df = df[~(df.isna()).any(axis=1)]

        # Get all columns starting with Rater
        rater_columns = [col for col in df.columns if col.startswith('Rater')]
        df['HumanRatersAgreeingCount'] = df.loc[:, df.columns.str.startswith("Rater")].apply(lambda x: x.value_counts().max(), axis=1)
        df['MajorityVote'] = df.loc[:, df.columns.str.startswith("Rater")].apply(lambda x: x.value_counts().idxmax(), axis=1)
        df["AtLeastTwoRatersAgree"] = df["HumanRatersAgreeingCount"] >= 2

        # Rule: When at least two humans agree
        total_frames_before = df.groupby(['Trial', 'Frame']).count().size
        total_trial_before = df.groupby(['Trial']).count().size

        df = df[df["AtLeastTwoRatersAgree"] == True]

        total_frames_toanalyse = df.groupby(['Trial', 'Frame']).count().size
        total_trial_toanalyse = df.groupby(['Trial']).count().size

        y_true = df.MajorityVote.values  # Gold standard is human
        y_pred = df.Algorithm.values

        # Filter out labels that are not present in y_true or y_pred
        unique_labels = np.unique(np.concatenate((y_true, y_pred)))
        labels = [label for label in desired_order if label in unique_labels]

        # Create a mapping for the filtered labels
        label_mapping = {label: i for i, label in enumerate(labels)}

        # Map y_true and y_pred to the new order
        mapped_y_true = np.array([label_mapping[label] for label in y_true if label in label_mapping])
        mapped_y_pred = np.array([label_mapping[label] for label in y_pred if label in label_mapping])

        # Plot the confusion matrix
        fig, ax = plt.subplots(1, 1, figsize=(7, 6))
        ConfusionMatrixDisplay.from_predictions(
            mapped_y_true, 
            mapped_y_pred, 
            ax=ax, 
            cmap=plt.cm.binary, 
            colorbar=False, 
            normalize='true',
            display_labels=labels,
            values_format='.2f',
            xticks_rotation='vertical'
        )

        true_counts = {label: np.sum(y_true == label) for label in labels}
        pred_counts = {label: np.sum(y_pred == label) for label in labels}

        # Create labels with true counts for y-axis and predicted counts for x-axis
        y_labels_with_true_counts = [f"{label} (N={true_counts[label]})" for label in labels]
        x_labels_with_pred_counts = [f"{label} (N={pred_counts[label]})" for label in labels]

        # Set the x and y tick labels with the new labels
        ax.set_xticklabels(x_labels_with_pred_counts, rotation='vertical')
        ax.set_yticklabels(y_labels_with_true_counts)

        ax.set_ylabel("Human rater annotation")
        ax.set_xlabel("Algorithmic annotation")

        fig.tight_layout()

        self.persistence.add_figure(SaveFigure(fig, "confusion_matrix"))